"""
Generates ChartVolt_Business_Plan.docx with embedded charts and flow schematics.
Run:  python Docs/_build_business_doc.py
"""

import os
import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(HERE, "_assets_business")
os.makedirs(ASSETS, exist_ok=True)

# Brand palette (electric "volt" theme)
NAVY = "#0B1F3A"
BLUE = "#1E6FE0"
CYAN = "#16C5D8"
AMBER = "#F5A623"
GREEN = "#27AE60"
SLATE = "#5B6B7F"
LIGHT = "#EAF1FB"
INK = "#1A2433"

plt.rcParams.update(
    {
        "font.family": "DejaVu Sans",
        "font.size": 11,
        "axes.edgecolor": SLATE,
        "axes.linewidth": 0.8,
        "figure.dpi": 200,
    }
)


def _box(ax, x, y, w, h, text, fc, tc="white", fs=11, bold=True):
    ax.add_patch(
        FancyBboxPatch(
            (x, y), w, h,
            boxstyle="round,pad=0.02,rounding_size=0.08",
            linewidth=0, facecolor=fc, mutation_aspect=1,
        )
    )
    ax.text(
        x + w / 2, y + h / 2, text,
        ha="center", va="center", color=tc, fontsize=fs,
        fontweight="bold" if bold else "normal", wrap=True,
    )


def _arrow(ax, x1, y1, x2, y2, color=SLATE):
    ax.add_patch(
        FancyArrowPatch(
            (x1, y1), (x2, y2),
            arrowstyle="-|>", mutation_scale=18,
            linewidth=2.2, color=color,
        )
    )


# ---------------------------------------------------------------- Flow of funds
def chart_flow_of_funds():
    fig, ax = plt.subplots(figsize=(9.2, 3.0))
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 30)
    ax.axis("off")

    y = 11
    h = 9
    boxes = [
        (1, "Card deposit\n(Nuvei / Stripe)", BLUE),
        (21, "Volt Credits\n100 = €1", CYAN),
        (41, "Play\nContests · 1v1 · Market", NAVY),
        (61, "Prizes\nto winners", GREEN),
        (81, "Withdrawal\nto bank", AMBER),
    ]
    w = 17
    for x, label, color in boxes:
        _box(ax, x, y, w, h, label, color, fs=10)
    for i in range(len(boxes) - 1):
        x1 = boxes[i][0] + w
        x2 = boxes[i + 1][0]
        _arrow(ax, x1 + 0.3, y + h / 2, x2 - 0.3, y + h / 2)

    # fee annotations
    notes = [
        (19.5, "+2% fee\n+VAT 21%"),
        (39.5, ""),
        (59.5, "−20% / −10%\nplatform fee"),
        (79.5, "−2% fee"),
    ]
    for x, t in notes:
        if t:
            ax.text(x, y - 1.5, t, ha="center", va="top", color=SLATE, fontsize=8.5)

    ax.text(50, 27.5, "How money moves through ChartVolt",
            ha="center", fontsize=13, fontweight="bold", color=INK)
    fig.tight_layout()
    p = os.path.join(ASSETS, "flow_of_funds.png")
    fig.savefig(p, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return p


# ---------------------------------------------------------------- Architecture
def chart_architecture():
    fig, ax = plt.subplots(figsize=(9.2, 4.8))
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 66)
    ax.axis("off")

    _box(ax, 35, 56, 30, 7, "Live market feed (Massive.com)", NAVY, fs=10)
    _box(ax, 33, 43, 34, 9,
         "PRIMARY SERVER\nprice engine · candle aggregation\nbackground jobs · writes", BLUE, fs=9.5)
    _box(ax, 73, 43.5, 23, 8, "MongoDB\nwallets · contests · ledger", "#3A4A5E", fs=9)
    _box(ax, 37, 29, 26, 7, "Redis relay  ·  pub/sub", SLATE, fs=9.5)
    _box(ax, 5, 15, 29, 9, "SECONDARY SERVER\nweb + real-time fan-out", CYAN, fs=9.5)
    _box(ax, 66, 15, 29, 9, "SECONDARY SERVER\nweb + real-time fan-out", CYAN, fs=9.5)
    _box(ax, 37, 2, 26, 8, "Users (web / mobile)", GREEN, fs=10)

    _arrow(ax, 50, 56, 50, 52.3)            # feed -> primary
    _arrow(ax, 67, 47.5, 73, 47.5)          # primary -> mongo
    _arrow(ax, 50, 43, 50, 36.3)            # primary -> redis
    _arrow(ax, 41, 29, 22, 24.3)            # redis -> left secondary
    _arrow(ax, 59, 29, 78, 24.3)            # redis -> right secondary
    _arrow(ax, 21, 15, 41, 10.3)            # left secondary -> users
    _arrow(ax, 80, 15, 59, 10.3)            # right secondary -> users

    ax.text(50, 65, "Built to scale: add servers as users grow",
            ha="center", fontsize=13, fontweight="bold", color=INK)
    fig.tight_layout()
    p = os.path.join(ASSETS, "architecture.png")
    fig.savefig(p, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return p


# ---------------------------------------------------------------- Revenue mix
def chart_revenue_mix():
    fig, ax = plt.subplots(figsize=(5.4, 4.4))
    labels = ["Contest &\nchallenge fees", "Deposit &\nwithdrawal fees",
              "Marketplace &\nGame Master", "Retained pools\n& creator fees"]
    sizes = [46, 28, 18, 8]
    colors = [BLUE, CYAN, AMBER, SLATE]
    wedges, _ = ax.pie(
        sizes, colors=colors, startangle=90,
        wedgeprops=dict(width=0.42, edgecolor="white", linewidth=2),
    )
    ax.legend(wedges, [f"{l}  ({s}%)" for l, s in zip(labels, sizes)],
              loc="center", fontsize=8.5, frameon=False,
              bbox_to_anchor=(0.5, -0.08), ncol=2)
    ax.text(0, 0, "Revenue\nmix*", ha="center", va="center",
            fontsize=12, fontweight="bold", color=INK)
    ax.set_title("Where revenue comes from (illustrative)",
                 fontsize=12, fontweight="bold", color=INK, pad=10)
    fig.tight_layout()
    p = os.path.join(ASSETS, "revenue_mix.png")
    fig.savefig(p, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return p


# ---------------------------------------------------------------- Forecast
def chart_forecast():
    fig, ax1 = plt.subplots(figsize=(7.4, 4.3))
    years = ["Year 1", "Year 2", "Year 3"]
    users = [2000, 10000, 30000]
    revenue = [0.19, 0.96, 2.9]  # € millions

    bars = ax1.bar(years, [u / 1000 for u in users], width=0.5,
                   color=LIGHT, edgecolor=BLUE, linewidth=1.5, label="Active users (000s)")
    ax1.set_ylabel("Avg. monthly active users (thousands)", color=SLATE, fontsize=10)
    ax1.tick_params(axis="y", labelcolor=SLATE)
    for b, u in zip(bars, users):
        ax1.text(b.get_x() + b.get_width() / 2, b.get_height() + 0.6,
                 f"{u:,}", ha="center", fontsize=9, color=SLATE)

    ax2 = ax1.twinx()
    ax2.plot(years, revenue, color=AMBER, marker="o", markersize=9,
             linewidth=3, label="Annual revenue (€M)")
    ax2.set_ylabel("Annual platform revenue (€ millions)", color=AMBER, fontsize=10)
    ax2.tick_params(axis="y", labelcolor=AMBER)
    for x, r in zip(years, revenue):
        ax2.annotate(f"€{r:.2f}M", (x, r), textcoords="offset points",
                     xytext=(0, 12), ha="center", fontsize=9.5,
                     fontweight="bold", color="#B9770E")
    ax2.set_ylim(0, 3.6)

    ax1.set_title("Illustrative 3-year trajectory",
                  fontsize=13, fontweight="bold", color=INK, pad=12)
    ax1.spines["top"].set_visible(False)
    ax2.spines["top"].set_visible(False)
    fig.tight_layout()
    p = os.path.join(ASSETS, "forecast.png")
    fig.savefig(p, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return p


# ---------------------------------------------------------------- DOCX helpers
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def style_base(doc):
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)
    st.font.color.rgb = RGBColor(0x1A, 0x24, 0x33)
    pf = st.paragraph_format
    pf.space_after = Pt(8)
    pf.line_spacing = 1.15


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "1E6FE0")
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1E, 0x6F, 0xE0)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p


def para(doc, text, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_image(doc, path, width=6.6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x5B, 0x6B, 0x7F)
    return p


def fee_table(doc, rows, headers):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Light List Accent 1"
    hdr = t.rows[0].cells
    for i, hh in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(hh)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], "0B1F3A")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t


# ---------------------------------------------------------------- Build doc
def build():
    flow = chart_flow_of_funds()
    arch = chart_architecture()
    mix = chart_revenue_mix()
    fc = chart_forecast()

    doc = Document()
    style_base(doc)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # --- Cover
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(90)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("ChartVolt")
    r.font.size = Pt(48)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Trade. Compete. Win.")
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x16, 0xC5, 0xD8)
    r.italic = True

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run("Business Plan & Company Overview")
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1E, 0x6F, 0xE0)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("Confidential · Prepared for internal and investor review")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x5B, 0x6B, 0x7F)
    doc.add_page_break()

    # --- A note up front (human voice)
    h1(doc, "A quick note before the detail")
    para(doc,
         "We built ChartVolt because we kept seeing the same thing happen to people who get "
         "into trading: they arrive excited, lose money fast, and quietly disappear. The markets "
         "are unforgiving, and most newcomers never get past that first painful month. We thought "
         "there had to be a better way to give people the thrill and the skill of trading without "
         "handing the markets their savings.")
    para(doc,
         "So we made trading a competition. You trade real live prices, but with virtual capital, "
         "and you compete for real prize pools. The most you can ever lose is a small entry fee you "
         "decide on up front. Everything that follows in this document explains how that idea works "
         "as a product, how it makes money, how we keep everyone safe, and where we plan to take it.")

    # --- 1 Product
    h1(doc, "1. What we are")
    para(doc,
         "ChartVolt is a gamified trading competition platform. The simplest way to describe it is "
         "\"fantasy sports for traders.\" People trade against real, live market prices, but the "
         "capital is virtual, so nobody is exposed to runaway losses. Instead they enter tournaments "
         "and one-on-one duels, climb leaderboards, and compete for prize pools funded by entry fees.")
    para(doc,
         "Underneath the game sits a genuine trading engine. Prices stream in real time, and players "
         "work with leverage, margin, stop-loss and take-profit orders, and liquidation, exactly as "
         "they would on a live platform. Today that engine is focused on forex, which is where the "
         "experience is sharpest and most reliable. Other asset classes are on the roadmap, and we "
         "are honest about that distinction throughout this document.")
    para(doc, "What players actually do on the platform:")
    bullet(doc, "Trade a realistic forex simulator driven by live market prices.")
    bullet(doc, "Enter competitions: multi-player tournaments with a prize pool and a live leaderboard.")
    bullet(doc, "Play 1v1 challenges: head-to-head duels where the winner takes the pot.")
    bullet(doc, "Spend Volt Credits in a marketplace for indicators, strategies, cosmetics, and creator packages.")
    bullet(doc, "Become a Game Master: host their own competitions and earn referral income.")
    bullet(doc, "Progress through levels, titles, badges, and a journey map that rewards them for sticking around.")
    para(doc,
         "A single wallet ties it together. Players deposit money, receive Volt Credits (our in-app "
         "currency, where 100 credits equals one euro), and use those credits to play, buy, and cash out.")

    # --- 2 Audience
    h1(doc, "2. Who it's for")
    para(doc,
         "Our core players are aspiring and retail traders who want to get better and compete, but "
         "are understandably nervous about risking real money on the markets. Right alongside them "
         "are competitive gamers who are drawn to ranks, prizes, and head-to-head play as much as to "
         "trading itself. The two groups overlap more than people expect.")
    para(doc,
         "Then there are the creators. Trading educators and community leaders are always looking for "
         "ways to engage their audiences, and our Game Master program lets them run branded contests "
         "and earn from the activity they generate. They are both customers and a distribution channel.")
    para(doc,
         "Finally, there is the business-to-business side. Because the whole platform is white-label, "
         "brokers, educators, and fintech brands can run ChartVolt under their own name. That turns a "
         "single consumer app into a product other companies pay to operate.")

    # --- 3 Opportunity
    h1(doc, "3. Why now")
    para(doc,
         "Three things are happening at once, and ChartVolt sits right where they meet. Retail trading "
         "has exploded, but the dropout rate is brutal because most people lose money early. At the "
         "same time, competitive and prize-based gaming has gone fully mainstream and proven that "
         "people will happily pay to compete. And there is a clear, unmet appetite for a way to enjoy "
         "trading without the open-ended financial risk.")
    para(doc,
         "Our angle is simple. We take an activity that normally costs people money and trust, and we "
         "turn it into something fun, social, repeatable, and capped in risk. That is better for the "
         "player and better for the business, because people stay longer and come back more often. "
         "And because the trading is simulated, we avoid much of the regulatory weight that comes with "
         "executing real trades.")

    # --- Flow of funds (schematic)
    h1(doc, "4. How the money flows")
    para(doc,
         "Money moves through the platform in a clean loop, and every step is recorded. A player "
         "deposits by card, receives credits, spends them on contests and items, wins prizes, and "
         "withdraws what they cash out. We earn a transparent fee at the deposit, the contest, and "
         "the withdrawal, never from a player's trading losses.")
    add_image(doc, flow, width=6.9)
    caption(doc, "Figure 1. The flow of funds, from card deposit to bank payout.")
    para(doc, "The fees behind that flow (these are our current defaults and are configurable per brand):")
    fee_table(
        doc,
        [
            ["Deposit fee", "2% on deposits"],
            ["Withdrawal fee", "2% on withdrawals"],
            ["Competition fee", "20% of the prize pool"],
            ["1v1 challenge fee", "10% of the prize pool"],
            ["Marketplace & Game Master sales", "the item / package price"],
            ["Unclaimed pools & lapsed creator fees", "retained by the platform"],
        ],
        ["Revenue source", "Default"],
    )
    para(doc,
         "Key figures, for reference: 100 Volt Credits equal one euro, the minimum deposit is €10, "
         "the minimum withdrawal is €20, VAT of 21% is applied where it is due, and Game Masters earn "
         "a 5% referral share by default. Behind the scenes we keep the player's wallet and the "
         "company's revenue on separate ledgers and reconcile them automatically, so the books always agree.")

    add_image(doc, mix, width=4.8)
    caption(doc, "Figure 2. Illustrative split of where revenue comes from. Actual mix depends on player behavior.")

    # --- 5 Protection
    h1(doc, "5. How we protect everyone")
    para(doc,
         "We treat safety as a feature, not a checkbox, and we think about it in three directions: "
         "protecting the company, protecting the payment channel and the bank, and protecting the "
         "customer. Each one has real machinery behind it.")
    h2(doc, "Protecting the company")
    para(doc,
         "We detect fraud and abuse with device fingerprinting, suspicion scoring, and multi-account "
         "and same-network checks that catch collusion and bonus abuse. Inside contests, anti-cheat "
         "rules guard against mirror-trading and abnormal drawdown, and we keep price snapshots so "
         "results can always be settled fairly. Every sensitive admin action is logged with its IP "
         "and device, and an automated reconciliation service constantly checks that every balance "
         "matches its transaction history.")
    h2(doc, "Protecting the bank and the payment channel")
    para(doc,
         "Every payment confirmation is cryptographically verified before we grant a single credit, "
         "so forged messages are rejected and logged. Deposits are processed in a way that can never "
         "double-count, even if the provider sends the same callback twice. When a chargeback comes "
         "in, the platform automatically assembles a professional evidence report, suspends the "
         "account, and reclaims the disputed credits. Deposits and withdrawals are rate-limited to "
         "shut down automated attacks.")
    h2(doc, "Protecting the customer")
    para(doc,
         "The biggest protection is the design itself: a player can only ever lose the small entry "
         "fee they chose, never an uncapped market loss. On top of that we run identity verification "
         "for withdrawals, enforce sensible withdrawal limits and hold periods, support two-factor "
         "authentication, and give every player a clear, itemized record of every euro and every "
         "credit, complete with proper invoices and VAT.")

    # --- 6 Architecture
    h1(doc, "6. Built to scale")
    para(doc,
         "None of this matters if the platform falls over when it gets busy, so we built it to grow. "
         "A primary server handles the live price feed, candle aggregation, and the heavy background "
         "work. Secondary servers handle the web traffic and push real-time updates to players, and "
         "they talk to the primary through a fast relay. When we need more capacity, we add another "
         "secondary server. There is no re-architecture required to scale up.")
    add_image(doc, arch, width=6.9)
    caption(doc, "Figure 3. The platform's multi-server design. Capacity grows by adding servers.")
    para(doc,
         "In plain terms: a couple of servers comfortably handle roughly a thousand active traders "
         "today, and the streaming side of the product scales horizontally well beyond that. The main "
         "thing we watch as we grow is the database write load on the primary, which is the natural "
         "ceiling and the first thing we would reinforce.")

    # --- 7 Marketing & Growth
    h1(doc, "7. How we grow")
    para(doc,
         "Our cheapest and most trustworthy growth comes from creators. The Game Master program turns "
         "influencers and community leaders into a distribution network: they host competitions for "
         "their followers and earn a referral share, so they bring the audience while we provide the "
         "platform. We complement that with focused performance marketing aimed at retail-trading and "
         "competitive-gaming audiences, and with flagship tournaments whose headline prize pools "
         "create their own publicity and viral leaderboards.")
    para(doc,
         "Keeping players is just as important as winning them. A low €10 entry point removes the "
         "barrier to trying, the progression systems give people a reason to come back daily, the "
         "leaderboards and duels create rivalry, and the built-in tutorials reduce the early confusion "
         "that usually drives newcomers away.")
    para(doc, "We grow in three phases:")
    bullet(doc, "Launch and prove it in one core market with forex contests and our first Game Masters.")
    bullet(doc, "Scale with creator-led and paid acquisition, more regions and languages, and a deeper marketplace.")
    bullet(doc, "Become a platform: sign white-label partners who each bring their own audience.")

    # --- 8 Forecast
    h1(doc, "8. The numbers")
    para(doc,
         "What follows is a planning model, not a promise. It shows how turnover behaves on a sensible "
         "set of assumptions, and the assumptions are meant to be argued with and replaced as we learn. "
         "The headline driver is the number of active players; everything else follows from that. We "
         "assume a blended revenue of roughly €8 per active player per month, drawn from the deposit, "
         "contest, and withdrawal fees described earlier.")
    add_image(doc, fc, width=6.6)
    caption(doc, "Figure 4. Illustrative three-year trajectory of active players and annual revenue.")
    fee_table(
        doc,
        [
            ["Year 1 (ramp)", "2,000", "~€190K"],
            ["Year 2 (scale)", "10,000", "~€960K"],
            ["Year 3 (platform + white-label)", "30,000", "~€2.9M"],
        ],
        ["Stage", "Avg. monthly active players", "Annual platform revenue"],
    )
    para(doc,
         "A note on what these figures mean: turnover here is the platform's own revenue from fees and "
         "sales, not the gross volume of money moving through player wallets, which is several times "
         "larger. The levers that move the model most are the number of active players, how much each "
         "one plays, how often they deposit, and how many white-label partners we sign. Our main "
         "variable cost is payment processing, which runs in the region of 2.9% plus a small fixed fee "
         "per deposit and is partly offset by our own deposit and withdrawal fees.")

    # --- Closing
    h1(doc, "In one breath")
    para(doc,
         "ChartVolt makes trading a fair, fun, capped-risk competition. Players trade live markets "
         "with virtual capital and compete for real prizes; we earn from transparent fees rather than "
         "from anyone's losses; we protect the company, the bank, and the customer at every layer; and "
         "we are built white-label on infrastructure that scales, so we can grow from a focused launch "
         "into the engine behind many branded trading-competition products.")

    out = os.path.join(HERE, "ChartVolt_Business_Plan.docx")
    doc.save(out)
    print("Saved:", out)


if __name__ == "__main__":
    build()
